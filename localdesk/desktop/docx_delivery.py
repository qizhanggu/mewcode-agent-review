"""DOCX 交付的生成、结构检查和实际渲染检查。

Markdown 仍是可审阅的源产物；DOCX 只能从已通过 Reviewer 的 Markdown
生成，先留在任务 staging 中。真实交付前必须同时具备结构检查和渲染
检查证据，不能把“文件后缀是 .docx”当成质量保证。
"""

from __future__ import annotations

import hashlib
import os
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Protocol

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
import fitz
from pypdf import PdfReader

from localdesk.desktop.workspace import DesktopWorkspace, WorkspaceError


class DocxDeliveryError(ValueError):
    pass


@dataclass(frozen=True)
class StructureCheck:
    approved: bool
    findings: list[str]
    paragraph_count: int
    heading_count: int
    source_count: int


@dataclass(frozen=True)
class RenderCheck:
    approved: bool
    findings: list[str]
    pdf_path: str | None
    image_paths: list[str]
    page_count: int


@dataclass(frozen=True)
class StagedDocx:
    staged_path: str
    final_path: str
    sha256: str
    source_markdown: str
    structure: StructureCheck
    render: RenderCheck
    summary: str


class DocxRenderer(Protocol):
    def render(self, docx_path: Path, output_dir: Path) -> RenderCheck: ...


class LibreOfficeDocxRenderer:
    """使用 LibreOffice 和 PyMuPDF 生成可检查的 PDF 与逐页 PNG。"""

    def __init__(self, soffice_path: str | None = None, pdftoppm_path: str | None = None, timeout_seconds: int = 90) -> None:
        self.soffice_path = soffice_path or _find_soffice()
        self.pdftoppm_path = pdftoppm_path or shutil.which("pdftoppm")
        self.timeout_seconds = timeout_seconds

    def render(self, docx_path: Path, output_dir: Path) -> RenderCheck:
        if not self.soffice_path:
            raise DocxDeliveryError("未找到 LibreOffice/soffice，无法完成 DOCX 实际渲染检查")
        output_dir.mkdir(parents=True, exist_ok=True)
        profile_dir = output_dir / "libreoffice-profile"
        profile_dir.mkdir(exist_ok=True)
        try:
            converted = subprocess.run(
                [
                    self.soffice_path,
                    "--headless",
                    f"-env:UserInstallation={profile_dir.as_uri()}",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_dir),
                    str(docx_path),
                ],
                capture_output=True,
                text=True,
                timeout=self.timeout_seconds,
                check=False,
            )
        except OSError as exc:
            raise DocxDeliveryError(f"启动 LibreOffice 失败: {exc}") from exc
        except subprocess.TimeoutExpired as exc:
            raise DocxDeliveryError("DOCX 渲染超时") from exc
        pdf_path = output_dir / f"{docx_path.stem}.pdf"
        if converted.returncode != 0 or not pdf_path.exists() or not pdf_path.stat().st_size:
            detail = (converted.stderr or converted.stdout).strip()
            raise DocxDeliveryError(f"LibreOffice 未生成有效 PDF: {detail or 'unknown error'}")
        try:
            pages = PdfReader(str(pdf_path)).pages
            page_count = len(pages)
            if page_count < 1 or any(not (page.extract_text() or "").strip() for page in pages):
                raise DocxDeliveryError("渲染后的 PDF 存在空白或不可提取文本页面")
        except DocxDeliveryError:
            raise
        except Exception as exc:
            raise DocxDeliveryError(f"无法检查渲染后的 PDF: {exc}") from exc
        images = _rasterize_pdf(pdf_path, output_dir)
        if len(images) != page_count or any(not image.stat().st_size for image in images):
            raise DocxDeliveryError("PDF 页面 PNG 渲染检查失败")
        return RenderCheck(True, [], str(pdf_path), [str(path) for path in images], page_count)


class FakeDocxRenderer:
    """离线测试渲染器；只验证 Runtime 编排，不替代真实 LibreOffice 验收。"""

    def render(self, docx_path: Path, output_dir: Path) -> RenderCheck:
        output_dir.mkdir(parents=True, exist_ok=True)
        image = output_dir / "page-1.png"
        image.write_bytes(b"\x89PNG\r\n\x1a\n")
        return RenderCheck(True, [], None, [str(image)], 1)


class DocumentDeliverySkill:
    """将经过审阅的 Markdown 转为可交付 DOCX。"""

    def __init__(self, workspace: DesktopWorkspace) -> None:
        self.workspace = workspace

    def stage_docx(self, task_id: str, source_markdown: str, filename: str, renderer: DocxRenderer) -> StagedDocx:
        source = Path(source_markdown)
        if not self.workspace.is_task_artifact(source) or not source.exists():
            raise WorkspaceError("DOCX 源 Markdown 必须是当前任务 staging 中已存在的草稿")
        name = _safe_docx_filename(filename)
        staged = self.workspace.task_dir(task_id) / "staging" / name
        final = self.workspace.output_root / name
        if final.exists():
            raise WorkspaceError(f"禁止覆盖已有文件: {final}")
        _markdown_to_docx(source, staged)
        structure = inspect_docx_structure(staged)
        if not structure.approved:
            raise DocxDeliveryError("DOCX 结构检查未通过: " + "; ".join(structure.findings))
        render = renderer.render(staged, self.workspace.task_dir(task_id) / "render" / staged.stem)
        if not render.approved:
            raise DocxDeliveryError("DOCX 渲染检查未通过: " + "; ".join(render.findings))
        return StagedDocx(
            staged_path=str(staged),
            final_path=str(final),
            sha256=_sha256(staged),
            source_markdown=str(source),
            structure=structure,
            render=render,
            summary="已通过结构与渲染检查的 DOCX 交付草稿",
        )


def inspect_docx_structure(path: Path) -> StructureCheck:
    findings: list[str] = []
    try:
        document = Document(path)
    except Exception as exc:
        return StructureCheck(False, [f"DOCX 无法打开: {exc}"], 0, 0, 0)
    paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text.strip()]
    headings = [paragraph for paragraph in paragraphs if paragraph.style and paragraph.style.name.startswith("Heading")]
    source_start = next((index for index, paragraph in enumerate(paragraphs) if paragraph.text.strip().lower() == "sources"), None)
    source_count = 0
    if not paragraphs:
        findings.append("DOCX 没有正文段落")
    if not headings:
        findings.append("DOCX 缺少结构化标题")
    if source_start is None:
        findings.append("DOCX 缺少 Sources 标题")
    else:
        source_count = sum(1 for paragraph in paragraphs[source_start + 1 :] if paragraph.text.strip())
        if source_count == 0:
            findings.append("DOCX Sources 下没有来源条目")
    if any("{{" in paragraph.text or "}}" in paragraph.text for paragraph in paragraphs):
        findings.append("DOCX 包含未替换模板占位符")
    return StructureCheck(not findings, findings, len(paragraphs), len(headings), source_count)


def _markdown_to_docx(source: Path, target: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Inches(1)
    section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    _configure_styles(document)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("LocalDesk Office Agent · staged document")
    _set_font(run, "Calibri", 8, "666666")

    for raw in source.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if match := re.match(r"^(#{1,3})\s+(.+)$", line):
            level = len(match.group(1))
            text = _clean_markdown(match.group(2))
            if level == 1:
                paragraph = document.add_paragraph()
                paragraph.style = document.styles["LocalDesk Title"]
                paragraph.add_run(text)
            else:
                document.add_heading(text, level=level - 1)
            continue
        if line.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(_clean_markdown(line[2:]))
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph(style="LocalDesk Note")
            paragraph.add_run(_clean_markdown(line[2:]))
            continue
        document.add_paragraph(_clean_markdown(line), style="Normal")
    target.parent.mkdir(parents=True, exist_ok=True)
    document.save(target)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (("Heading 1", 16, "2E74B5", 16, 8), ("Heading 2", 13, "2E74B5", 12, 6)):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    title = document.styles.add_style("LocalDesk Title", WD_STYLE_TYPE.PARAGRAPH)
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(10)
    note = document.styles.add_style("LocalDesk Note", WD_STYLE_TYPE.PARAGRAPH)
    note.font.name = "Calibri"
    note.font.size = Pt(10)
    note.font.color.rgb = RGBColor.from_string("555555")
    note.paragraph_format.space_after = Pt(8)


def _set_font(run: object, name: str, size: int, color: str) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)


def _clean_markdown(text: str) -> str:
    return text.replace("`", "").replace("**", "").replace("__", "")


def _safe_docx_filename(filename: str) -> str:
    name = Path(filename).name
    if name != filename or not name or name in {".", ".."}:
        raise WorkspaceError("DOCX 文件名必须是不含目录的文件名")
    if not name.lower().endswith(".docx"):
        name += ".docx"
    if not re.fullmatch(r"[\w.\-\u4e00-\u9fff ]+", name):
        raise WorkspaceError("DOCX 文件名包含不允许的字符")
    return name


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _find_soffice() -> str | None:
    candidates = [
        os.getenv("LOCALDESK_SOFFICE_PATH"),
        shutil.which("soffice"),
        shutil.which("libreoffice"),
        r"D:\Apps\LibreOffice\program\soffice.exe",
        r"C:\Program Files\LibreOffice\program\soffice.exe",
    ]
    return next((candidate for candidate in candidates if candidate and Path(candidate).is_file()), None)


def _rasterize_pdf(pdf_path: Path, output_dir: Path) -> list[Path]:
    try:
        pdf = fitz.open(pdf_path)
        images: list[Path] = []
        for number, page in enumerate(pdf, start=1):
            image = output_dir / f"page-{number}.png"
            page.get_pixmap(dpi=150, alpha=False).save(image)
            images.append(image)
        return images
    except Exception as exc:
        raise DocxDeliveryError(f"无法将渲染后的 PDF 栅格化为 PNG: {exc}") from exc
